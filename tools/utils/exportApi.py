#!/usr/bin/env python
# -*- coding:utf-8 -*-
# ====#====#====#====
# Author:CXY
# CreateDate: 2022/9/15 10:24
# Filename:exportApi.py
# Function:
# ====#====#====#====
import datetime
import sys
import time

import re
import docx
import uuid
import pymysql
from docx.document import Document
from docx.table import _Cell, Table
from docx.oxml.table import CT_Tbl
from docx.oxml.text.paragraph import CT_P
from docx.text.paragraph import Paragraph
file_path = '/Users/liujin/Desktop/脚本/导出接口/Piufoto.docx'
mysql_config = dict(host='1.117.81.152', user='liuma', passwd='YcSewyTdBeM8SWLz')
file = docx.Document(file_path)  # 导出的api文件
# database = 'api_copy'  # 需要写入的数据库，api为真实数据库，api_copy为测试数据库
database = 'api'  # 需要写入的数据库，api为真实数据库，api_copy为测试数据库
pause_index = 0  # 从什么地方开始执行
raise_index = sys.maxsize


class OpenMysql:
    def __init__(self, *args, **kwargs):
        """
        kwargs like host='www', user='root', passwd='123456', db='test', port=3306
        """
        # 将传进来的变量保存到self,不能在这个函数进行conn的创建
        # 因为初始化变量后不一定会执行变量的__exit__,容易造成僵尸连接
        self.mysql_config = kwargs

    def __enter__(self):
        try:
            import pymysql
        except Exception:
            raise Exception('pymysql must be installed at your environment')
        self.conn = pymysql.connect(**self.mysql_config)
        self.cursor = self.conn.cursor()
        return self

    def execute(self, sql):
        self.ping()
        self.cursor.execute(sql)
        result = self.cursor.fetchall()
        return result

    def commit(self):
        self.ping()
        self.conn.commit()

    def ping(self):
        try:
            self.conn.ping()
        except:
            self.__enter__()

    def __exit__(self, exc_type, exc_val, exc_tb):
        self.cursor.close()
        self.conn.close()
        del self


# 数据库中域名对应的uuid
domain_map = {
    'm': 'be73459f-9585-42db-8373-602255d05b97',
    'ct': '1ed6e70d-0a6c-431a-8b65-96758547c6c0',
    'nm': '9dcdf3ac-59e6-4ba9-b66d-612181f7d6be',
    'www': 'e28cd6a7-6e1c-40a1-a960-a1cbe717b44d',
    'am.': '60f04ac3-340c-4dba-b804-b2366a1b71e5',
    'cs': 'ed07c436-9cce-4317-97f6-308e967ce62b',
    'rps': 'd37a22a3-e350-4ee7-94ca-af325eed676d',
    '': '',
    'myHost': '',
    'ma4': '91cb3dd3-8c56-4749-834e-8160e905af1a',
    'v4c': '09af25ce-9134-4012-8287-9ed4840f1b6f',
    'k': 'cf635e39-3231-4edb-bb12-e6eee48ff538'
}


# 接口类
class Api_case:
    def __init__(self, description='', url='', request_method='', body='{"raw": "", "file": [], "form": [], '
                                                                       '"json": "", "type": "form-urlencoded"}',
                 domain='',
                 protocol='', name='', form_list=None):
        if form_list is None:
            form_list = []
        self.description = description
        self.url = url
        self.request_method = request_method
        self.body = body
        self.domain = domain
        self.protocol = protocol
        self.name = name
        self.form_list = form_list


# 统计数据库中的接口数量
def statistics_case_num(mysqlc):
    sql = 'select count(*) from liuma.{database};'.format(database=database)
    data = mysqlc.execute(sql)
    return data[0][0]


def analyze_case(tmp_case: Api_case):
    """
    分析数据
    """
    # 对末尾的数据进行处理，去除加密的参数
    for re_i in ['/v{validParams}', '/{{signature}}']:
        tmp_case.url = tmp_case.url.replace(re_i, '')
    # 如果接口以https开头，就将protocol设置为HTTPS，去除接口中的HTTP前缀
    if tmp_case.url.startswith('https://'):
        tmp_case.protocol = 'HTTPS'
        tmp_case.url = tmp_case.url[8:]
    elif tmp_case.url.startswith('http://'):
        tmp_case.protocol = 'HTTP'
        tmp_case.url = tmp_case.url[7:]
    else:
        tmp_case.protocol = 'HTTPS'
    # 将域名写入case中
    if tmp_case.domain != '':
        tmp_case.domain = tmp_case.domain.replace('/alltuu', '')
    else:
        tmp_case.domain, tmp_case.url = tmp_case.url.split('/', 1)
        tmp_case.url = '/' + tmp_case.url
        tmp_case.url=tmp_case.url.replace(' ', '')
    tmp_case.domain = re.sub('.(?:alltuu|guituu).com', '', tmp_case.domain)
    # 将请求方式改为大写
    tmp_case.request_method = tmp_case.request_method.upper()
    # 如果接口名没有的话，就使用接口路径代替
    if not tmp_case.name:
        tmp_case.name = str(tmp_case.url)
    if tmp_case.form_list != []:
        tmp_case.body = '{{"raw": "", "file": [], "form": {form}, "json": "", "type": "form-urlencoded"}}'.format(
            form=str(tmp_case.form_list).replace("'", "\"").replace('True', 'true').replace('False', 'false'))
    return tmp_case


def transfer_data_type(param_type):
    if param_type.upper() == 'STRING':
        return 'String'
    elif param_type.upper() == 'INT':
        return 'String'
    elif param_type.upper() == 'FLOAT':
        return 'String'
    elif param_type.upper() == 'Float':
        return 'String'
    elif param_type.upper() == 'BOOLEAN':
        return 'Boolean'
    elif param_type.upper() == 'JSONOBJECT':
        return 'JSONObject'
    elif param_type.upper() == 'JSONARRAY':
        return 'JSONArray'
    elif param_type.upper() == 'FILE':
        return 'File'
    else:
        return 'String'


def transfer_require(require):
    if require == '是':
        return True
    else:
        return False


def interface_parameters(child, index, tmp_Api_case: Api_case):
    """处理接口参数"""

    if file.paragraphs[index].text in (
            '请求Body参数', 'json字段说明', '请求参数', '新增字段', '二期-付费下载', '参数'):
        # 请求参数示例有些问题的，备注就先不理他了，应为是请求参数
        # 不加人签名的参数
        table = Table(child, file)
        print(file.paragraphs[index].text)
        try:
            index_param = 0
            index_required = 1
            index_type = 2
            index_description = 3
            for index_row, row in enumerate(table.rows):
                # 如果是第一行，就是参数名，识别出每一行的参数名代表什么意思，然后写入对应位置
                if index_row == 0:
                    for index_col, col in enumerate(row.cells):
                        if '参数' in col.text:
                            index_param = index_col
                        elif '类型' in col.text:
                            index_type = index_col
                        elif '说明' in col.text:
                            index_description = index_col
                        elif '必选' in col.text:
                            index_required = index_col
                else:
                    param_type = transfer_data_type(row.cells[index_type].text)
                    required = transfer_require(row.cells[index_required].text)
                    tmp_Api_case.form_list.append(
                        {"name": row.cells[index_param].text,
                         "type": param_type,
                         "value": '',
                         "required": required
                         }
                    )
            return False
        except:
            print('表格找的有问题', table.cell(0, 0).text)


def interface_details(index, tmp_Api_case):
    """解析接口信息"""
    # print('第{index}段的内容是：{context}'.format(index=index,context=file.paragraphs[index].text))
    # 如果是标题一 todo 需要将标题一的API作为一组API分类
    # 遇到标题二或三后开始
    api_item_begin = False
    if file.paragraphs[index].style.name.startswith('Heading 1'):
        # file.paragraphs[index].text 从数据中找到是否一致的组，然后没有的话就写入数据库中，有的话，读取出uuid
        pass
    # 如果是标题二或标题三，将里面的数据读取出来，按数据要求整理数据，写入数据库中
    elif file.paragraphs[index].style.name.startswith('Heading 2') or file.paragraphs[index].style.name.startswith(
            'Heading 3'):
        # 本来想在这里做一次初始化的，但是有些接口文档中，存在两个标题三，如果初始化会出问题
        # tmp_Api_case = Api_case()
        re_pat = '([\u4e00-\u9fa5]|[a-zA-Z]).*'
        try:
            tmp_Api_case.name = re.search(re_pat, file.paragraphs[index].text).group()
        except:
            print("Error")
            tmp_Api_case.name = 'Normal'
            pass
        api_item_begin = True
        # 对数据进行处理
    elif file.paragraphs[index].style.name.startswith('Heading 5'):
        if file.paragraphs[index].text == '简要描述':
            tmp_Api_case.description = file.paragraphs[index + 1].text
        elif file.paragraphs[index].text in ('请求URL', '请求路径'):
            tmp_Api_case.url = file.paragraphs[index + 1].text
        elif file.paragraphs[index].text == '请求方式':
            tmp_Api_case.request_method = file.paragraphs[index + 1].text
        elif file.paragraphs[index].text in ('请求域名', '域名'):
            tmp_Api_case.domain = file.paragraphs[index + 1].text
        elif file.paragraphs[index].text in ('成功返回示例', '返回示例', '备注'):
            return False, api_item_begin
    return True, api_item_begin


def write_database(case, mysqlc):
    """将数据写入数据库"""
    count = 10000 + int(statistics_case_num(mysqlc))
    sql = "insert into liuma.{database} (id, name, level, module_id, project_id, method, path, protocol, domain_sign, description, header, body, query, rest,  create_user, update_user, create_time, update_time, status) values " \
          "('{uuid}', '{name}','p1','0e5cc7da-9e2e-4e77-adee-6dced6074464','3333f9b1-b81a-11ed-8641-52540085c262','{method}','{path}','{protocol}','{domain_sign}','{description}','[]','{body}','[]','[]','py','py',{time},{time},'Normal')" \
        .format(uuid=str(uuid.uuid1()), name=case.name, method=case.request_method, path=case.url,
                description=case.description, protocol=case.protocol,
                domain_sign=domain_map[re.sub('[^a-zA-Z.0-9]', '', case.domain)],
                time=int(time.time() * 1000), body=case.body, database=database)
    count += 1
    # 执行sql语句
    mysqlc.execute(sql)
    # 提交到数据库执行
    mysqlc.commit()


# 开始执行
def run(mysqlc):
    if isinstance(file, Document):
        parent_elm = file.element.body
    elif isinstance(file, _Cell):
        parent_elm = file._tc
    else:
        raise ValueError("something's not right")

    index = -1
    item_begin = True
    api_item_begin = False
    tmp_Api_case = Api_case()
    for child in parent_elm.iterchildren():
        # 只读取100行的数据先
        if index > raise_index:
            raise Exception("运行手动结束")

        if isinstance(child, CT_Tbl):
            # 如果表格是请求参数的形式，将它写进来，不然不管他
            if index < pause_index:
                continue
            item_begin = interface_parameters(child, index, tmp_Api_case)
        elif isinstance(child, CT_P):
            # 如果是文字段落的话，进行解析
            index += 1
            if index < pause_index:
                continue
            item_begin, tmp_flag = interface_details(index, tmp_Api_case)
            api_item_begin = api_item_begin or tmp_flag
        #  如果接口文档结束了，并且符合开始条件，就开始处理case数据，并写到数据库中
        if not item_begin and index >= pause_index and api_item_begin:
            # 对提取出来的数据进行处理
            case = analyze_case(tmp_Api_case)
            # 写入数据库
            print(index, case.name)
            write_database(case, mysqlc)

            # 一个接口结束，初始化数据
            tmp_Api_case = Api_case()
            item_begin = True
            api_item_begin = False


if __name__ == '__main__':
    with OpenMysql(**mysql_config) as mysqlc:
        run(mysqlc)